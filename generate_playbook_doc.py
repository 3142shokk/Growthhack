#!/usr/bin/env python3
"""Generate Higgsfield Counter-Playbook — professional Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

OUT = "/Users/alemzhanakhmetzhanov/Desktop/projects/growthhack/docs/Higgsfield_Counter_Playbook.docx"

# ── colour palette ────────────────────────────────────────────────────
BLACK      = RGBColor(0x1a, 0x1a, 0x1a)
NAVY       = RGBColor(0x1e, 0x27, 0x4d)   # section headings
INDIGO     = RGBColor(0x3d, 0x52, 0xa0)   # sub-headings / accent
MID        = RGBColor(0x44, 0x44, 0x55)   # body text
LIGHT      = RGBColor(0x6b, 0x6b, 0x80)   # captions / labels
RULE       = RGBColor(0xcc, 0xcc, 0xd8)   # horizontal rules
WHITE      = RGBColor(0xff, 0xff, 0xff)
COVER_BG   = RGBColor(0x1e, 0x27, 0x4d)

doc = Document()

# ── page margins ─────────────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(1.1)
    sec.top_margin  = sec.bottom_margin = Inches(0.9)

# ── default paragraph font ────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(5)

# ─────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), color="CCCCCC", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para(doc, text="", size=10.5, bold=False, italic=False,
         color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=5, left_indent=0, line_spacing=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if line_spacing:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(line_spacing)
    if text:
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(size)
        run.bold       = bold
        run.italic     = italic
        run.font.color.rgb = color if color else BLACK
    return p

def add_run(p, text, size=10.5, bold=False, italic=False, color=None):
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    run.font.color.rgb = color if color else BLACK
    return run

def h1(doc, text):
    p = para(doc, space_before=18, space_after=6)
    add_run(p, text.upper(), size=13, bold=True, color=NAVY)
    rule(doc, color="1e274d", thickness="12")
    return p

def h2(doc, text):
    p = para(doc, space_before=12, space_after=4)
    add_run(p, text, size=11, bold=True, color=INDIGO)
    return p

def h3(doc, text):
    p = para(doc, space_before=6, space_after=3)
    add_run(p, text, size=10.5, bold=True, color=MID)
    return p

def body(doc, text, color=None, size=10.5, italic=False, indent=0):
    return para(doc, text, size=size, color=color or BLACK,
                italic=italic, space_after=4, left_indent=indent)

def caption(doc, text, indent=0):
    return para(doc, text, size=9, color=LIGHT, italic=True,
                space_after=3, left_indent=indent)

def bullet_item(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10.5)
    run.font.color.rgb = color or BLACK
    return p

def rule(doc, color="CCCCCC", thickness="6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def spacer(doc, pts=6):
    p = para(doc, space_before=0, space_after=pts)
    return p

def kv(doc, key, value, key_color=None, val_color=None, indent=0):
    p = para(doc, space_after=3, left_indent=indent)
    add_run(p, key + ":  ", size=10, bold=True, color=key_color or LIGHT)
    add_run(p, value,       size=10.5, color=val_color or BLACK)
    return p

def label_block(doc, label, text, indent=0.1):
    caption(doc, label.upper(), indent=indent)
    body(doc, text, indent=indent + 0.05)

def template_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name   = "Courier New"
    run.font.size   = Pt(9.5)
    run.font.color.rgb = MID
    # shaded background via paragraph border trick
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F7")
    pPr.append(shd)
    return p

def simple_table(doc, headers, rows, col_widths=None, header_bg="1e274d"):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"
    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name  = "Calibri"
        run.font.size  = Pt(9.5)
        run.bold       = True
        run.font.color.rgb = WHITE
    # data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F5F5FA"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name  = "Calibri"
            run.font.size  = Pt(9.5)
            run.font.color.rgb = BLACK
    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    spacer(doc, 8)
    return tbl


# ═════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════

p = para(doc, "GROWTH ENGINEERING TRACK  |  HACKNU 2026",
         size=9, color=INDIGO, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=36, space_after=4, bold=True)

p = para(doc, "Higgsfield Counter-Playbook",
         size=30, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=6)

rule(doc, color="3d52a0", thickness="16")

p = para(doc, "Replicating Claude's Viral Growth Mechanics for a Competing AI Product",
         size=13, italic=True, color=MID,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=24)

p = para(doc, "Part 4 of 4  |  Based on 197,190 scraped data points",
         size=10, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p = para(doc, "Reddit (84,088 posts)  |  Hacker News (107,000 stories)  |  YouTube  |  X / Twitter",
         size=10, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p = para(doc, "r/HiggsfieldAI (n=2,094)  |  Archetype corpus (n=74,000 posts)",
         size=10, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

rule(doc, color="CCCCCC")
spacer(doc, 4)
p = para(doc, "Confidential  |  HackNU 2026",
         size=9, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═════════════════════════════════════════════════════════════════════
h1(doc, "Contents")
toc = [
    ("01", "Executive Summary",           "3"),
    ("02", "The 3 Mechanisms to Replicate","4"),
    ("03", "Audience Targeting",           "6"),
    ("04", "Channel Strategy",             "8"),
    ("05", "Content Format Playbook",      "10"),
    ("06", "Amplifier Recruitment",        "15"),
    ("07", "Narrative Moment Calendar",    "16"),
    ("08", "8-Week Execution Calendar",    "17"),
    ("09", "Metrics That Matter",          "19"),
    ("10", "Data Trail",                   "21"),
]
for num, title, pg in toc:
    p = para(doc, space_after=3)
    add_run(p, num + "  ", size=10, bold=True, color=INDIGO)
    add_run(p, title, size=10.5, color=BLACK)
    add_run(p, "   " + pg, size=10, color=LIGHT)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 01  EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════
h1(doc, "01  Executive Summary")

body(doc, (
    "Claude's viral growth is not luck. It runs on three identifiable, repeatable mechanics that "
    "can be directly applied to Higgsfield AI. This playbook documents those mechanics, explains "
    "why each works from the data, and translates each into a concrete execution plan. Every "
    "recommendation is traceable to a specific finding from the scraped corpus."
))
spacer(doc, 4)
h3(doc, "The three core mechanisms:")

mech_rows = [
    ("M1  Threat > Launch",
     "Claude's threat disclosure got 33.6M views. Best product launch: 5M. 6.7x gap. "
     "The Hollywood AI disruption narrative is Higgsfield's equivalent opportunity."),
    ("M2  Insider Leak Archetype",
     "230 Insider Leak posts avg 102 score vs 1,208 Capability Demo posts avg 24. "
     "4.25x ROI gap. Higgsfield is almost entirely Showcase content (avg 4.6 score)."),
    ("M3  Copy Formula",
     "'just' as opening word = 638x viral lift. 'Higgsfield' as opener = 1.26x. "
     "Title sweet spot: 70-130 characters. None of Higgsfield's current posts apply this."),
]
simple_table(doc,
    headers=["Mechanism", "Finding"],
    rows=mech_rows,
    col_widths=[2.0, 4.2])

body(doc, (
    "This playbook applies each mechanism to Higgsfield through four practical lenses: "
    "audience targeting, channel strategy, content formats, and an 8-week execution calendar. "
    "A metrics section defines exactly what success looks like and what action each signal requires."
))
doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 02  THE 3 MECHANISMS
# ═════════════════════════════════════════════════════════════════════
h1(doc, "02  The 3 Mechanisms to Replicate")
body(doc, (
    "Claude's viral growth runs on identifiable patterns. The three below were selected because "
    "they are (a) consistently present across platforms, (b) directly applicable to Higgsfield's "
    "context, and (c) grounded in statistically significant sample sizes from the scraped corpus."
))
spacer(doc, 6)

mechanisms = [
    {
        "title": "M1  Threat Over Launch",
        "claude_data": "Distillation attack disclosure: 33.6M views. Best product launch: 5M views. Gap: 6.7x.",
        "why": (
            "Claude's highest-traffic content is not about what Claude is — it is about what Claude "
            "means to the world. The threat disclosure worked because 'state-level actors are "
            "targeting us' communicates scale more viscerally than any benchmark or feature release. "
            "Every major Reddit and HN spike follows a real-world event, not a product launch calendar."
        ),
        "for_hf": (
            "The Hollywood AI displacement conversation — WGA, SAG-AFTRA, studio policy, actor "
            "likeness rights — is live, urgent, and covered by mainstream press. It is the "
            "equivalent of Claude's DeepSeek and national-security moments. A Higgsfield post with "
            "something credible to say inside that conversation gets surfaced to an audience far "
            "larger than Higgsfield's current follower base. The opportunity requires preparation: "
            "positions must be drafted before the news breaks, not after."
        ),
        "source": "Finding 1: threat-over-launch  |  Top 20 viral moments table  |  Weekly engagement timeline",
    },
    {
        "title": "M2  Insider Leak Archetype Dominance",
        "claude_data": "Insider Leak / Reveal: n=230, avg score 102. Capability Demo: n=1,208, avg score 24. Gap: 4.25x.",
        "why": (
            "The community produces 20x more Capability Demos than Insider Leaks, yet Insider Leaks "
            "return 4x higher engagement per post. Scarcity is structurally load-bearing: scarcity "
            "is maintained because genuine access and timing advantage are rare. The moment this "
            "archetype becomes saturated its score converges toward the mean. The ROI exists "
            "precisely because most brands never produce it."
        ),
        "for_hf": (
            "r/HiggsfieldAI is 39% Showcase posts averaging 4.6 score. That is the Capability Demo "
            "equivalent — the most-produced, lowest-ROI archetype. One genuine Insider post — a "
            "working filmmaker, a real production outcome including honest failures, a specific cost "
            "comparison — is worth approximately 20 Showcase posts in total engagement."
        ),
        "source": "Finding 2: archetype-trap  |  Reddit corpus n=74,000  |  r/HiggsfieldAI flair analysis n=2,094",
    },
    {
        "title": "M3  Copy Formula",
        "claude_data": (
            "Opening word 'just': 638x viral lift. Opening word 'claude': 1.26x. "
            "Title length sweet spot: 70-130 characters."
        ),
        "why": (
            "Performing opening words are story hooks, not product claims. The formula behind the "
            "top-performing posts: [recency word] + [what happened] + [specific surprising outcome], "
            "70-130 characters. 'Just had Sonnet rewrite my entire test suite in 40 minutes' — "
            "recency, personal, version-specific, outcome, time frame. Posts that open with the "
            "product name perform 1.26x above baseline. Posts that open with 'just' perform 638x "
            "above baseline. The difference is story hook vs. product announcement."
        ),
        "for_hf": (
            "'Higgsfield releases cinematic AI' scores 1.26x. "
            "'Just tried to recreate a Spielberg tracking shot with AI for $0' scores potentially 638x. "
            "Every Higgsfield post title should pass one test before publishing: "
            "does it start with a story hook or a product claim?"
        ),
        "source": "Finding 3: copy-formula  |  Reddit title analysis n=74,000  |  Opening word lift chart (6c)",
    },
]

for m in mechanisms:
    h2(doc, m["title"])
    caption(doc, "Source: " + m["source"])
    kv(doc, "Claude data", m["claude_data"])
    label_block(doc, "Why it works", m["why"])
    label_block(doc, "Application for Higgsfield", m["for_hf"])
    spacer(doc, 6)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 03  AUDIENCE TARGETING
# ═════════════════════════════════════════════════════════════════════
h1(doc, "03  Audience Targeting")
body(doc, (
    "Claude's amplifier stack has eight or more independent voices, each serving a different "
    "distribution function. Higgsfield currently has one (EHuanglu, responsible for 69.4% of "
    "community X views). The targeting strategy below defines four segments and the specific "
    "approach to reach and activate each."
))
spacer(doc, 6)

segments = [
    {
        "id": "S1", "name": "Working Film and VFX Professionals",
        "size": "50,000 - 200,000 active on Reddit and HN",
        "channels": "r/VFX (150k members), r/Filmmakers (1.5M members), r/HiggsfieldAI, HN Show HN",
        "content": "F1 Filmmaker Insider Reveal, F2 Production Breakdown, F4 Constraint / Workaround",
        "recruit": (
            "Direct offer of early access combined with a specific project brief: 'We will fund one "
            "shot for your next project using Higgsfield.' Not a generic trial offer — a concrete, "
            "funded use case relevant to their existing work."
        ),
        "why": (
            "This is the Karpathy / rubenhassid equivalent — Claude's researcher amplifiers whose "
            "technical posts reach audiences that treat them as credible peers. Those six accounts "
            "averaged 1.8 million views per post. Film professionals carry equivalent trust weight "
            "in the creative industry."
        ),
    },
    {
        "id": "S2", "name": "Creator and AI Hobbyist Community",
        "size": "2,094 posts currently in r/HiggsfieldAI",
        "channels": "r/HiggsfieldAI (owned, primary), r/generativeAI, r/aivideo",
        "content": "F3 Competitor Comparison, F4 Constraint / Workaround, weekly community challenges",
        "recruit": (
            "Offer moderator status and early feature access to the top ten contributors by post "
            "count. BholaCoder (112 posts), memerwala_londa (74 posts, avg score 9.8), and "
            "la_dehram (48 posts, avg score 14.3) are already producing consistent volume. "
            "A community role converts them from neutral users to advocates."
        ),
        "why": (
            "These users are currently posting KLING content in r/HiggsfieldAI because there is no "
            "content strategy holding them on-topic. KLING posts in the subreddit average 6.6 score "
            "vs 2.6 for Higgsfield's own video model posts. Activating this segment closes the "
            "competitor infiltration gap."
        ),
    },
    {
        "id": "S3", "name": "Technical and Developer Audience",
        "size": "0 Higgsfield HN posts vs 107,000 for Claude",
        "channels": "Hacker News (Show HN), r/MachineLearning, r/MediaSynthesis",
        "content": "F2 Production Breakdown in HN format, technical architecture writeups",
        "recruit": (
            "Direct message to adocomplete (10 stories, avg 947 HN points) and meetpateltech "
            "(27 stories, avg 368 points) with a one-sentence pitch. These are enthusiasts who "
            "surface things they find genuinely interesting to engineers. One front-page placement "
            "creates developer credibility infrastructure that persists for years."
        ),
        "why": (
            "HN is Claude's single largest platform at 107,000 posts and 58% of total content "
            "volume. HN front-page placement drives tech press pickup within 24-48 hours and "
            "generates engineering Slack and developer Discord shares. Zero Higgsfield presence "
            "means zero developer credibility."
        ),
    },
    {
        "id": "S4", "name": "Industry Press and Policy Layer",
        "size": "Variety, The Hollywood Reporter, The Verge, WIRED",
        "channels": "Hacker News (credibility primer), X (news cycle), LinkedIn",
        "content": "F5 Industry Stance, F6 Narrative Surfing posts",
        "recruit": (
            "Become quotable before press coverage is needed. Post a clear policy position, a "
            "real cost-data comparison, a real filmmaker outcome with numbers. Journalists covering "
            "Hollywood AI need sources who have something specific and verifiable to say. "
            "Higgsfield is positioned to be that source."
        ),
        "why": (
            "Claude's Dario Amodei statement on AI and national security received 18.5 million "
            "views — the second highest viral moment in the dataset — because it was a credible "
            "position inside an ongoing mainstream news story. Higgsfield needs one equivalent "
            "moment inside the Hollywood AI policy conversation."
        ),
    },
]

for s in segments:
    h2(doc, s["id"] + "  " + s["name"])
    kv(doc, "Estimated reach", s["size"])
    kv(doc, "Primary channels", s["channels"])
    kv(doc, "Content types", s["content"])
    label_block(doc, "Recruitment strategy", s["recruit"])
    label_block(doc, "Why this segment", s["why"])
    spacer(doc, 6)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 04  CHANNEL STRATEGY
# ═════════════════════════════════════════════════════════════════════
h1(doc, "04  Channel Strategy")
body(doc, (
    "Five channels ranked by ROI. Owned channels take priority over earned and borrowed. "
    "All posting frequency and timing recommendations are derived from Claude engagement "
    "data and the r/HiggsfieldAI corpus."
))
spacer(doc, 6)

channels = [
    {
        "name": "r/HiggsfieldAI  (Owned — Primary)",
        "type_freq": "Owned  |  3 to 5 posts per week",
        "current": (
            "2,094 posts, avg score 4.8. KLING-flaired posts: 153 (avg score 6.6). "
            "Higgsfield video model posts: 46 (avg score 2.6). Competitors are outperforming "
            "Higgsfield's own model in Higgsfield's own community."
        ),
        "tactics": [
            "Audit flair policy: rename competitor model flairs (KLING, SORA, WAN) to 'AI Video Comparison'. The goal is redirect, not ban — competitor comparisons are high-engagement when Higgsfield is the primary subject.",
            "Create an 'HF Insider' flair for team posts and early-access creator content. Announcement posts already average 13.2 — the highest flair score in the subreddit. This is the archetype to feed.",
            "Pin a Weekly Challenge thread every Monday. This gives the creator community a recurring on-topic prompt and drives consistent scheduled volume.",
            "Apply the copy formula to every post: 70-130 characters, story hook opener ('just', 'when', 'i'm'). Benchmark target: beat the 4.8 subreddit average.",
            "Track weekly: average score, competitor content percentage (target below 3%), Showcase vs Insider and Tutorial ratio.",
        ],
    },
    {
        "name": "YouTube  (Owned — High Priority)",
        "type_freq": "Owned  |  2 posts per week, Tuesday and Thursday",
        "current": (
            "Higgsfield YouTube engagement rate: 5.12% vs Claude's 3.02%. "
            "Content quality is demonstrably strong. The problem is discovery — "
            "not enough people are finding it."
        ),
        "tactics": [
            "Produce 10 SEO-optimized tutorials within 60 days. Title targets: 'AI filmmaking workflow for indie directors', 'how to maintain character consistency in AI video', 'replace green screen with AI step-by-step'.",
            "Each tutorial must include a specific limitation and workaround section. Workaround content outperforms showcase content in the Reddit data (avg 16.8 vs 15.7) and drives search discovery.",
            "Cross-promote to r/HiggsfieldAI on the same day using a personal-story post about making the video. Use the story hook title formula, not the YouTube title.",
            "Upload on Tuesday and Thursday. These are peak engagement days from Claude's timeline data.",
        ],
    },
    {
        "name": "Hacker News  (Earned — High Priority)",
        "type_freq": "Earned  |  1 submission per month",
        "current": "0 Higgsfield submissions vs 107,000 Claude posts. Complete developer credibility gap.",
        "tactics": [
            "First submission: 'Show HN: How we maintain consistent actor faces across AI video shots without per-person fine-tuning'. Technical, specific, invites expert discussion.",
            "Second submission (Month 2): open-source evaluation framework for AI video quality — give engineers a tool they can actually use.",
            "Submit on Tuesday or Wednesday, 9 to 11 AM Eastern. Never submit on Fridays or weekends.",
            "Direct message adocomplete and meetpateltech the day before each submission with a one-sentence pitch. These accounts are enthusiasts, not gatekeepers.",
            "Every HN post must include: technical architecture, real numbers, honest trade-offs, and an open question at the end to invite comments.",
        ],
    },
    {
        "name": "X / Twitter  (Earned — Medium Priority)",
        "type_freq": "Earned  |  1 official post per day plus amplifier network",
        "current": (
            "EHuanglu generates 69.4% of Higgsfield community views. "
            "Single-point-of-failure risk is high."
        ),
        "tactics": [
            "Official account: post at Tuesday through Thursday, 14:00 to 19:00 UTC. This is the peak engagement window from Claude's timeline analysis.",
            "Format mix: 70% video clips with specific outcome text, 30% text threads with data or findings. Open every post with 'just', 'when', or 'i'm' — not the brand name.",
            "Amplifier strategy: identify 5 film and VFX creators with 100,000 to 500,000 followers. Target: reduce EHuanglu dependency from 69.4% to below 30% within 6 months.",
            "Maintain three pre-drafted narrative surfing posts (see Section 07) ready for deployment within 4 hours of a Hollywood AI news trigger.",
        ],
    },
    {
        "name": "Shared Subreddits  (Borrowed — Low Priority)",
        "type_freq": "Borrowed  |  2 to 3 posts per week from top-performing owned content",
        "current": (
            "r/generativeAI, r/aivideo, r/Filmmakers. Higgsfield competes directly with "
            "Runway and Sora for the same feed position."
        ),
        "tactics": [
            "Cross-post only after content has performed in r/HiggsfieldAI with a score of 10 or above. Never as the first destination.",
            "Modify titles for shared subreddits: remove or soften the Higgsfield brand name in the opener. Lead with the outcome or the problem.",
            "r/Filmmakers (1.5 million members) is the highest-ROI shared subreddit for F1 and F5 format content — film professionals respond to insider and stance archetypes.",
            "Track the owned-to-shared ratio weekly. If shared sub posts consistently outperform owned posts, the content may be too brand-focused for the owned community.",
        ],
    },
]

for ch in channels:
    h2(doc, ch["name"])
    p = para(doc, space_after=3)
    add_run(p, ch["type_freq"], size=10, italic=True, color=LIGHT)
    label_block(doc, "Current state", ch["current"])
    caption(doc, "TACTICS", indent=0.1)
    for t in ch["tactics"]:
        bullet_item(doc, t)
    spacer(doc, 6)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 05  CONTENT FORMAT PLAYBOOK
# ═════════════════════════════════════════════════════════════════════
h1(doc, "05  Content Format Playbook")
body(doc, (
    "Six formats derived from Claude's archetype and title-analysis data. "
    "Every template, title formula, and platform assignment is derived from "
    "what actually performed, not from general best-practice guidance."
))
caption(doc, (
    "Copy rules applied to all formats: open with 'just', 'when', or 'i'm' — not the brand name  |  "
    "title length 70-130 characters  |  state the specific outcome in the title"
))
spacer(doc, 8)

formats = [
    {
        "number": "F1", "name": "Filmmaker Insider Reveal",
        "avg": "102 avg Reddit score",
        "archetype": "Insider Leak / Reveal",
        "source": "n=230 posts, avg 102. 4.25x better than Capability Demo (avg 24). Scarcest high-ROI archetype in the dataset.",
        "why": (
            "The highest-scoring archetype in the dataset. Rare because it requires genuine access "
            "or timing advantage. Most brands never produce it. The element that makes it perform "
            "is the honest failure — posts framed as personal discoveries with real trade-offs "
            "outperform polished showcases because they are useful to readers facing the same decisions."
        ),
        "template": (
            "I'm a [professional role] with [N] years in [field]. We used Higgsfield on [real project type].\n"
            "Here is what actually happened:\n\n"
            "What worked: [specific capability and time or cost saved]\n"
            "What broke: [specific failure with context]\n"
            "Cost vs. traditional production: [real numbers]\n"
            "Would I use it again: [honest verdict with conditions]"
        ),
        "formula": "70-130 chars  |  'I'm a [credential]' or 'just used Higgsfield on [real context]'  |  state the honest outcome",
        "example": "I'm a VFX supervisor. We used AI video for the first time on a paid shoot. Honest breakdown — what it saved and what it broke.",
        "platforms": "r/HiggsfieldAI (first), r/VFX, r/Filmmakers",
        "targeting": "S1 — Film and VFX Professionals",
    },
    {
        "number": "F2", "name": "Production Breakdown — HN Format",
        "avg": "242 avg HN points",
        "archetype": "Model Release / Technical",
        "source": "HN Model Release archetype: avg 242 pts, n=1,187 stories. Highest-performing HN archetype in the dataset.",
        "why": (
            "Technical depth is what HN rewards. An architecture writeup, not a demo. The HN audience "
            "shares to engineering Slack channels, developer Discords, and technical newsletters. "
            "Claude's highest-scoring HN content explains HOW something works, invites critique, "
            "and treats the reader as an engineer. The same content posted to Reddit uses a "
            "personal-story frame with the same underlying information."
        ),
        "template": (
            "Show HN: [specific technical problem we solved]\n\n"
            "Why this problem is hard: []\n"
            "Our approach: [actual method, no marketing language]\n"
            "Results: [specific, measurable outcomes]\n"
            "Trade-offs: [what you lose, what breaks at scale]\n"
            "Open question: [genuine unresolved problem — invites expert comments]"
        ),
        "formula": "HN: 'Show HN: How we [solved specific technical problem]'  |  Reddit: 'just figured out [specific capability]'",
        "example": "Show HN: How we maintain consistent actor identity across AI video shots without per-person fine-tuning",
        "platforms": "HN (Show HN), r/MachineLearning, r/HiggsfieldAI",
        "targeting": "S3 — Technical and Developer Audience",
    },
    {
        "number": "F3", "name": "Competitor Comparison",
        "avg": "+56% score lift",
        "archetype": "Controversy / Data-Driven",
        "source": "Reddit competitor-mention posts avg 28 vs 18 for solo-brand posts (56% lift, n=74,000 posts). 'kling' already appears in Higgsfield's top-content wordcloud.",
        "why": (
            "The r/HiggsfieldAI community is already posting KLING content (153 posts, avg score 6.6). "
            "A structured, data-backed Higgsfield vs Kling comparison posted by Higgsfield would "
            "surface in both communities simultaneously, score above the competitor-mention baseline, "
            "and take ownership of a comparison the community is already conducting organically."
        ),
        "template": (
            "Tested Higgsfield vs [Kling / Runway / Sora] on [specific use case]: [verdict in title]\n\n"
            "Test setup: same prompt, both tools, [N] runs\n"
            "Evaluation criteria: [specific metrics]\n"
            "Result 1: [specific outcome with numbers]\n"
            "Result 2: [specific outcome with numbers]\n"
            "Where Higgsfield wins: [honest assessment]\n"
            "Where it loses: [honest assessment]\n"
            "Conclusion: [verdict with use-case specificity]"
        ),
        "formula": "Include competitor name explicitly  |  include specific test methodology  |  state the verdict in the title",
        "example": "Higgsfield vs Kling on cinematic action sequences: I ran 50 prompts. Full breakdown with scores.",
        "platforms": "r/HiggsfieldAI, r/aivideo, r/generativeAI",
        "targeting": "S2 (Creator Community) and S1 (Film Professionals)",
    },
    {
        "number": "F4", "name": "Constraint and Workaround",
        "avg": "16.8 avg Reddit score",
        "archetype": "Workaround",
        "source": "Workaround flair avg 16.8 vs Built_with_Claude avg 15.7. 'limit' is one of the highest-frequency words in the top-500 post wordcloud.",
        "why": (
            "The honest limitation post outperforms the showcase because it is useful. Readers have "
            "already hit the same wall and are searching for a solution. This format generates "
            "long-tail search traffic because people with the same problem search for it specifically. "
            "Claude's community generates workaround content organically — it is the second "
            "most-discoverable content type in the corpus. Higgsfield has almost none."
        ),
        "template": (
            "Higgsfield [specific limitation] — here is the workaround\n\n"
            "Problem: [what breaks, why it breaks, what you were trying to do]\n\n"
            "Workaround:\n"
            "Step 1: [concrete action]\n"
            "Step 2: [concrete action]\n"
            "Step 3: [concrete action]\n\n"
            "Result: [before and after — specific, not vague]\n"
            "Trade-off: [what you lose with this method]"
        ),
        "formula": "Name the constraint explicitly  |  promise the solution  |  be honest about the trade-off  |  title is searchable",
        "example": "Higgsfield struggles with long-take tracking shots — here is how I get consistent motion with 3-second clips and a transition technique.",
        "platforms": "r/HiggsfieldAI, r/aivideo",
        "targeting": "S2 — Creator Community (long-tail search traffic)",
    },
    {
        "number": "F5", "name": "Industry Stance",
        "avg": "102 avg Reddit score",
        "archetype": "Controversy / Stance",
        "source": "Controversy / Stance archetype ties Insider Reveal at avg 102 (n=266 posts). Dario Amodei DoW statement: 18.5M views. Pete Hegseth response: 17.6M views.",
        "why": (
            "The entertainment industry is in an active dispute about AI. WGA and SAG-AFTRA have "
            "live AI policies. Trade press covers it weekly. Claude rode the policy conversation "
            "to 17.6 million views by having something specific and credible to say. Higgsfield is "
            "the only AI video company that could contribute from a filmmaking perspective rather "
            "than from an AI company perspective."
        ),
        "template": (
            "[Clear, specific position on AI and filmmaking]\n\n"
            "Context: I have used Higgsfield on [real project context].\n"
            "What I actually found: [specific data point or observation]\n\n"
            "The claim I am pushing back on: [specific]\n"
            "Why it is more nuanced:\n"
            "  - [Specific point 1]\n"
            "  - [Specific point 2 with evidence from real experience]\n\n"
            "The question that actually matters is not [common framing].\n"
            "It is [better, more specific framing]."
        ),
        "formula": "State the position directly  |  'I think [controversial claim]' works  |  'Some perspectives on...' does not  |  BREAKING framing has 2.7x emotional trigger lift",
        "example": "AI video is not replacing cinematographers. It is replacing the $500-per-day rental equipment decisions. Here is why that is a completely different problem.",
        "platforms": "r/Filmmakers, r/movies, HN, X (during news spikes)",
        "targeting": "S4 (Industry Press) and S1 (Film Professionals)",
    },
    {
        "number": "F6", "name": "Narrative Surfing Post",
        "avg": "9.3x view boost at peak (DeepSeek lesson)",
        "archetype": "Third-Party Validation / Controversy",
        "source": "DeepSeek mentions gave Claude tweets 9.3x view boost (2.8M vs 0.3M baseline) during Jan-Feb 2026. The mechanism is algorithmic surfacing into an active news feed.",
        "why": (
            "When a major news story is active, content that references it is surfaced to everyone "
            "following that story — audiences far larger than Higgsfield's current follower base. "
            "The window is narrow: the DeepSeek effect disappeared within weeks of the news cycle "
            "ending. Preparation must happen before the trigger, not after."
        ),
        "template": (
            "[React to specific news event — deploy within 2 to 4 hours]\n\n"
            "Here is what [event] means from an actual filmmaker's perspective:\n\n"
            "[1-2 sentences of honest reaction with specificity]\n\n"
            "What I can tell you from using AI video tools for [N] months:\n"
            "  - [Specific data point from Higgsfield experience]\n"
            "  - [Honest limitation relevant to the policy debate]\n\n"
            "The conversation should be about [reframe that positions Higgsfield credibly]."
        ),
        "formula": "Reference the news event directly  |  lead with filmmaker perspective  |  draft BEFORE the news breaks — window is 2-4 hours",
        "example": "SAG-AFTRA just updated their AI video policy. I have used AI video tools for 6 months. Here is what the policy gets right and what it misses.",
        "platforms": "X (primary, within 2 hours), r/Filmmakers, r/movies, HN if technical angle",
        "targeting": "All segments plus new audience carried by the news cycle",
    },
]

for f in formats:
    h2(doc, f"{f['number']}  {f['name']}  —  {f['avg']}")
    caption(doc, "Source: " + f["source"])
    kv(doc, "Archetype", f["archetype"])
    label_block(doc, "Why it works", f["why"])
    h3(doc, "Post Template")
    template_block(doc, f["template"])
    kv(doc, "Title formula", f["formula"])
    kv(doc, "Example title", f["example"])
    kv(doc, "Post to", f["platforms"])
    kv(doc, "Targeting", f["targeting"])
    spacer(doc, 8)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 06  AMPLIFIER RECRUITMENT
# ═════════════════════════════════════════════════════════════════════
h1(doc, "06  Amplifier Recruitment")
body(doc, (
    "EHuanglu generates 69.4% of Higgsfield's community X views. Claude has eight or more "
    "independent amplifier voices, each serving a distinct distribution function. Every amplifier "
    "added reduces single-point-of-failure risk and raises the view floor."
))
spacer(doc, 6)

h2(doc, "Claude's Amplifier Model")
caption(doc, "Reference structure being replicated")
simple_table(doc,
    headers=["Role", "Account", "Function"],
    rows=[
        ("Depth engine",          "AnthropicAI",           "580M total views, 923 tweets, official narrative anchor"),
        ("Technical credibility", "karpathy / rubenhassid", "6-8 posts each, avg 1.8M+ views, reaches developer and researcher audiences"),
        ("Cross-industry signal", "SenSanders / andrewchen","Non-AI audiences — signals cultural significance beyond tech"),
        ("Rapid surface",         "ns123abc",               "17 tweets, avg 2.4M views — finds and surfaces breaking stories fast"),
    ],
    col_widths=[1.5, 1.8, 3.0])

spacer(doc, 4)
h2(doc, "Higgsfield Target Amplifier Stack")
caption(doc, "Build within 8 weeks. Each role is an independent insurance policy.")
simple_table(doc,
    headers=["Role", "Criteria", "Recruitment approach", "Timeline"],
    rows=[
        ("Technical filmmaker",      "DoP or VFX supervisor, 50k-200k followers, posts technical creative content",         "Early access and funded shot for their next project",                            "Week 5-6"),
        ("Indie director / creator", "YouTube filmmaker, 100k-500k subscribers, regularly reviews AI tools",                "Exclusive feature preview 2 weeks before public launch",                        "Week 5-6"),
        ("AI video enthusiast",      "Already in r/HiggsfieldAI, 20k-100k followers, posts weekly",                        "Moderator role and direct line to team for bug reports",                        "Week 1-2"),
        ("Industry insider",         "Film school educator, trade journalist, or film festival programmer",                  "Data from the corpus: real cost and workflow comparisons they can cite as source", "Week 7-8"),
        ("Cross-industry voice",     "Non-film creator who uses video for business (course creator, event videographer)",    "Free credits and use-case brief for their specific workflow",                    "Week 7-8"),
    ],
    col_widths=[1.4, 2.0, 2.2, 0.7])

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 07  NARRATIVE MOMENT CALENDAR
# ═════════════════════════════════════════════════════════════════════
h1(doc, "07  Narrative Moment Calendar")
body(doc, (
    "Claude's number one viral moment was a threat disclosure with 33.6 million views. "
    "DeepSeek gave Claude a 9.3x X view boost by being in the right conversation at the right time. "
    "The Hollywood AI story is Higgsfield's equivalent. All three triggers below must be "
    "monitored continuously and response content must be pre-drafted before any trigger fires."
))
spacer(doc, 6)

triggers = [
    {
        "name": "WGA / SAG-AFTRA AI Policy Update",
        "probability": "High — active negotiations with regular updates",
        "sources": "Variety, The Hollywood Reporter, Deadline Hollywood  |  r/movies  |  r/Screenwriting",
        "window": "2-4 hours after publication",
        "boost": "3 to 10x based on DeepSeek mechanism (9.3x at peak news cycle)",
        "angles": [
            "What AI video policy should actually require from AI companies — from the perspective of someone who uses Higgsfield on real productions",
            "The real cost comparison: AI video vs. traditional production for independent vs. studio-level projects",
            "What fair attribution looks like for AI-assisted filmmaking, with Higgsfield's approach as a concrete example",
        ],
    },
    {
        "name": "Major AI Video Competitor Launches or Fails Publicly",
        "probability": "Medium-high — Sora, Runway, and Pika release updates regularly",
        "sources": "TechCrunch, The Verge, r/AIVideo  |  X trending",
        "window": "Within 4 hours of announcement",
        "boost": "56% score lift from competitor mention plus news cycle amplification",
        "angles": [
            "Direct comparison: specific capability tested side-by-side with real outputs",
            "What the new model gets right that Higgsfield does not (honest assessment earns credibility)",
            "The specific capability where Higgsfield demonstrably outperforms, with supporting evidence",
        ],
    },
    {
        "name": "Major Hollywood Studio Announces AI Policy or AI-Made Production",
        "probability": "Medium — Netflix, Disney, and A24 all have in-progress AI policies",
        "sources": "The Hollywood Reporter, Variety, Wall Street Journal  |  LinkedIn industry posts",
        "window": "Same day — LinkedIn and X simultaneously, Reddit the following day",
        "boost": "High reach from industry audiences plus potential press pickup via HN credibility primer",
        "angles": [
            "A filmmaker's view on what this means for independent productions, not studio budgets",
            "The specific workflow this policy enables or restricts that was not possible before",
            "What responsible disclosure looks like for AI-assisted content — Higgsfield's stated position",
        ],
    },
]

for t in triggers:
    h2(doc, t["name"])
    kv(doc, "Probability",         t["probability"])
    kv(doc, "Monitor sources",     t["sources"])
    kv(doc, "Deployment window",   t["window"])
    kv(doc, "Expected boost",      t["boost"])
    h3(doc, "Pre-drafted angles (write now, publish when triggered)")
    for angle in t["angles"]:
        bullet_item(doc, angle)
    spacer(doc, 6)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 08  8-WEEK EXECUTION CALENDAR
# ═════════════════════════════════════════════════════════════════════
h1(doc, "08  8-Week Execution Calendar")
body(doc, (
    "Foundation before volume. Volume before amplification. Measurement at Week 8. "
    "Any format averaging below 8 score in r/HiggsfieldAI after Week 8 is discontinued — "
    "the data determines the format mix for Sprint 2, not the plan."
))
spacer(doc, 6)

calendar = [
    {
        "week": "Week 1-2", "phase": "Foundation",
        "focus": "Fix infrastructure before adding volume",
        "actions": [
            "Audit r/HiggsfieldAI flair policy. Rename KLING, SORA, and WAN competitor flairs to 'AI Video Comparison'. KLING currently averages 6.6 score vs Higgsfield video model at 2.6 — competitors are outperforming the brand in its own community.",
            "Post 10 foundation pieces: 4x Filmmaker Story (F1), 3x Constraint and Workaround (F4), 2x Production Breakdown (F2), 1x Competitor Comparison (F3). Apply title formula to all.",
            "Apply the copy formula to every post: 70-130 characters, story hook opener ('just', 'when', 'i'm'). Benchmark: beat the 4.8 subreddit average.",
            "Offer moderator status to the top three r/HiggsfieldAI contributors by post count: BholaCoder (112 posts), memerwala_londa (74 posts, avg 9.8), la_dehram (48 posts, avg 14.3).",
            "Draft three pre-written narrative surfing posts (F6) for the three most likely news triggers from Section 07. Store as unpublished drafts.",
        ],
    },
    {
        "week": "Week 3-4", "phase": "Content Volume",
        "focus": "Establish posting rhythm, launch YouTube, first HN submission",
        "actions": [
            "Post 3x per week to r/HiggsfieldAI. Format mix: 2x Constraint and Workaround (F4), 1x Competitor Comparison (F3). Track score against the 4.8 baseline.",
            "Launch first YouTube tutorial: 'AI filmmaking workflow for indie directors'. SEO-optimized title, specific outcome, workaround section included. Upload on Tuesday.",
            "Submit first Show HN: 'Show HN: How we maintain consistent actor identity across AI video shots without per-person fine-tuning'. Direct message adocomplete the day before with a one-sentence pitch.",
            "Research 5 film and VFX professionals for the amplifier recruitment pipeline. Note their specific use cases and content style. Do not reach out yet.",
            "Cross-post top-performing r/HiggsfieldAI posts (score 10 or above) to r/aivideo and r/generativeAI with modified titles.",
        ],
    },
    {
        "week": "Week 5-6", "phase": "Amplification",
        "focus": "Activate amplifiers, second YouTube video, narrative monitoring",
        "actions": [
            "Direct message adocomplete and meetpateltech with the second HN story pitch: 'We built an open-source quality evaluation framework for AI video — here is the repo and findings.'",
            "Outreach to 3 film and VFX professionals: early access plus a specific project brief funded by Higgsfield. Not a generic trial offer — a concrete production use case.",
            "Upload second YouTube tutorial: 'How to maintain character consistency in AI video — step-by-step'. Upload on Thursday.",
            "Post one F5 Industry Stance piece if any Hollywood AI news event has occurred in the previous two weeks.",
            "Review r/HiggsfieldAI average score. If above 4.8, the archetype shift is producing results. If not, audit the format mix of posts published in Weeks 3-4.",
        ],
    },
    {
        "week": "Week 7-8", "phase": "First Measurement",
        "focus": "Measure, cut what failed, plan Sprint 2",
        "actions": [
            "Pull all metrics: r/HiggsfieldAI avg score (target 7+), competitor content percentage (target below 5%), YouTube avg views (target above 20,000), HN points earned (target above 100).",
            "Classify the top 5 posts by score: what archetype were they? The format mix for Sprint 2 follows the data, not the original plan.",
            "Discontinue any format averaging below 8 score in r/HiggsfieldAI. No exceptions.",
            "Amplifier check: have any outreach contacts posted about Higgsfield? If not, send a follow-up with a fresh angle and an exclusive asset.",
            "HN check: if the first submission did not reach front page, analyze what did front-page that week and adjust the second submission accordingly.",
        ],
    },
]

for phase in calendar:
    h2(doc, f"{phase['week']}  —  {phase['phase']}")
    caption(doc, "Focus: " + phase["focus"])
    for action in phase["actions"]:
        bullet_item(doc, action)
    spacer(doc, 6)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 09  METRICS THAT MATTER
# ═════════════════════════════════════════════════════════════════════
h1(doc, "09  Metrics That Matter")
body(doc, (
    "Six metrics derived directly from the scraped data. Each has a current baseline from "
    "the corpus, a Week 8 milestone, and a long-term target. Each alert specifies the action "
    "required — not merely the condition that triggered it."
))
spacer(doc, 6)

metrics = [
    {
        "name": "r/HiggsfieldAI Average Score",
        "formula": "Mean score of all r/HiggsfieldAI posts in rolling 30-day window",
        "now": "4.8", "week8": "8+", "target": "15+",
        "alert": (
            "Below 4.8 sustained across three or more weeks: the archetype mix is wrong. "
            "Check the Showcase percentage (currently 39%, avg score 4.6) against Announcement, "
            "Tutorial, and Insider content (avg 13.2). Produce more Announcement-format content immediately."
        ),
    },
    {
        "name": "Competitor Content Ratio",
        "formula": "Competitor-flaired posts (KLING + SORA + WAN) divided by total r/HiggsfieldAI posts",
        "now": "~8%", "week8": "<4%", "target": "<2%",
        "alert": (
            "Above 10%: the subreddit is functioning as a neutral AI video forum, not a Higgsfield "
            "community. Enforce flair policy. KLING scoring 6.6 vs Higgsfield video model at 2.6 "
            "is a community capture problem requiring immediate moderation action."
        ),
    },
    {
        "name": "Archetype Distribution",
        "formula": "Percentage of posts classified as Insider Reveal or Controversy / Stance in trailing 30 days",
        "now": "~0%", "week8": "10%+", "target": "20%+",
        "alert": (
            "Below 10% for three consecutive weeks: the content team is defaulting to Showcase format. "
            "Run an Insider Reveal briefing — assign a specific filmmaker to the format with a "
            "concrete project brief and a target post date."
        ),
    },
    {
        "name": "HN Points Earned",
        "formula": "Total HN points from Higgsfield-originated or Higgsfield-seeded submissions",
        "now": "0", "week8": "50+", "target": "500+",
        "alert": (
            "Zero HN points after four weeks: content is not technical enough for the platform. "
            "A Show HN requires architecture and honest trade-offs, not a demo. "
            "Zero front-page stories after two submissions: direct message adocomplete with the "
            "next submission pitch before it goes live."
        ),
    },
    {
        "name": "YouTube Average Views",
        "formula": "Average views per video published in trailing 30 days",
        "now": "33,000", "week8": "40,000+", "target": "75,000+",
        "alert": (
            "Engagement rate stays above 5% but views are not growing after Week 4: this is an "
            "SEO and discovery problem, not a content quality problem. Audit video titles against "
            "YouTube search volume for target keywords. "
            "Engagement rate drops below 4%: content quality is regressing — return to specific "
            "outcome format with workaround section."
        ),
    },
    {
        "name": "Independent Amplifier Count",
        "formula": "Number of X accounts with 50,000+ followers posting positively about Higgsfield in trailing 30 days",
        "now": "1", "week8": "3", "target": "5+",
        "alert": (
            "Still 1 after Week 6: outreach from Weeks 5-6 has not converted. Send a follow-up "
            "with a fresh angle and an exclusive asset — early access to an unreleased feature, "
            "a behind-the-scenes technical writeup, or a co-creation brief. "
            "Reaching 5 independent voices reduces EHuanglu dependency below 30%."
        ),
    },
]

simple_table(doc,
    headers=["Metric", "Now", "Week 8", "Target"],
    rows=[(m["name"], m["now"], m["week8"], m["target"]) for m in metrics],
    col_widths=[3.2, 0.7, 0.8, 0.8])

spacer(doc, 4)
for m in metrics:
    h2(doc, m["name"])
    caption(doc, "Formula: " + m["formula"])
    label_block(doc, "Alert and required action", m["alert"])
    spacer(doc, 4)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 10  DATA TRAIL
# ═════════════════════════════════════════════════════════════════════
h1(doc, "10  Data Trail")
body(doc, (
    "Every recommendation in this document is traceable to a specific measurement from the "
    "scraped corpus. The table below maps each key claim to its source data."
))
spacer(doc, 6)

trail = [
    ("Insider Reveal avg 102",                                       "Reddit corpus, archetype classifier, n=230 posts"),
    ("Capability Demo avg 24",                                       "Reddit corpus, archetype classifier, n=1,208 posts"),
    ("Threat disclosure 33.6M vs product launch 5M (6.7x gap)",     "Top 20 viral moments table, Anthropic accounts + r/ClaudeAI"),
    ("Competitor mention +56% score lift",                           "Reddit corpus, 2,098 competitor posts vs 71,914 solo-brand posts"),
    ("Opening word 'just' = 638x viral lift",                        "Reddit title analysis, top-10% viral posts vs baseline (n=74,000)"),
    ("Title length sweet spot 70-130 characters",                    "Reddit score vs title length regression, n=74,000 posts"),
    ("Workaround flair avg 16.8",                                    "Reddit flair engagement chart, n=1,937 workaround posts"),
    ("HN Model Release avg 242 points",                              "HN corpus, archetype breakdown, n=1,187 posts"),
    ("adocomplete: 10 stories, avg 947 HN points",                   "HN top submitters chart, selectivity vs volume analysis"),
    ("Higgsfield YouTube engagement 5.12% vs Claude 3.02%",          "YouTube scrape, Higgsfield channel (likes + comments divided by views)"),
    ("r/HiggsfieldAI: 2,094 posts, avg 4.8, KLING 153 posts avg 6.6 vs HF video model 46 posts avg 2.6",
                                                                     "r/HiggsfieldAI scrape (n=2,094), flair breakdown and score analysis"),
    ("EHuanglu = 69.4% of Higgsfield community X views",             "X amplifier scrape, official vs community split chart (cmp_07)"),
    ("DeepSeek 9.3x X view boost",                                   "X competitor timeline Jan-Feb 2026, Anthropic account scrape"),
    ("Announcement flair avg 13.2 (highest in r/HiggsfieldAI)",      "r/HiggsfieldAI flair analysis, n=33 announcement posts"),
    ("Claude Reddit 11.5x growth in 12 months",                      "r/ClaudeAI monthly post volume chart (cmp_02)"),
    ("Dario Amodei DoW statement 18.5M views",                       "Top 20 viral moments table, Anthropic X account scrape"),
    ("BREAKING emotional trigger: 2.7x lift",                        "Reddit emotional trigger analysis, top-10% viral posts (chart 6a)"),
    ("Humor flair avg 79.4",                                          "Reddit flair engagement chart, n=2,521 humor posts"),
]

simple_table(doc,
    headers=["Claim", "Source"],
    rows=trail,
    col_widths=[3.0, 3.3])

doc.add_page_break()

# ── save ─────────────────────────────────────────────────────────────
import os
os.makedirs("/Users/alemzhanakhmetzhanov/Desktop/projects/growthhack/docs", exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
